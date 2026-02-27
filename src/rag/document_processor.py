"""
文档处理器 - 支持多种文档格式解析
支持格式: PDF, Word(.docx), Markdown(.md), 纯文本(.txt)
"""

import os
from pathlib import Path
from typing import List, Dict, Any, Optional
from dataclasses import dataclass
import logging

# 文档解析库
from pypdf import PdfReader
from docx import Document

logger = logging.getLogger(__name__)


@dataclass
class Document:
    """文档数据结构"""
    content: str
    metadata: Dict[str, Any]
    source: str
    doc_type: str


class DocumentProcessor:
    """
    多格式文档处理器
    
    支持的格式:
    - PDF: 使用pypdf解析
    - Word: 使用python-docx解析
    - Markdown: 纯文本读取
    - TXT: 纯文本读取
    """
    
    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.md', '.txt'}
    
    def __init__(self, encoding: str = 'utf-8'):
        """
        初始化文档处理器
        
        Args:
            encoding: 文本编码，默认utf-8
        """
        self.encoding = encoding
    
    def process_file(self, file_path: str) -> Document:
        """
        处理单个文件
        
        Args:
            file_path: 文件路径
            
        Returns:
            Document对象
            
        Raises:
            ValueError: 不支持的文件格式
            FileNotFoundError: 文件不存在
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        ext = path.suffix.lower()
        
        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"不支持的文件格式: {ext}. "
                f"支持的格式: {', '.join(self.SUPPORTED_EXTENSIONS)}"
            )
        
        # 根据文件类型选择解析方法
        parsers = {
            '.pdf': self._parse_pdf,
            '.docx': self._parse_docx,
            '.md': self._parse_text,
            '.txt': self._parse_text,
        }
        
        content = parsers[ext](path)
        
        # 构建元数据
        metadata = {
            'filename': path.name,
            'extension': ext,
            'size_bytes': path.stat().st_size,
            'modified_time': path.stat().st_mtime,
        }
        
        doc_type = ext.lstrip('.')
        
        logger.info(f"成功解析文档: {path.name}, 类型: {doc_type}, 字符数: {len(content)}")
        
        return Document(
            content=content,
            metadata=metadata,
            source=str(path),
            doc_type=doc_type,
        )
    
    def process_directory(
        self, 
        dir_path: str, 
        recursive: bool = True
    ) -> List[Document]:
        """
        处理目录下的所有支持格式的文件
        
        Args:
            dir_path: 目录路径
            recursive: 是否递归处理子目录
            
        Returns:
            Document对象列表
        """
        path = Path(dir_path)
        
        if not path.exists():
            raise FileNotFoundError(f"目录不存在: {dir_path}")
        
        if not path.is_dir():
            raise ValueError(f"不是目录: {dir_path}")
        
        documents = []
        
        # 获取文件列表
        if recursive:
            files = path.rglob('*')
        else:
            files = path.glob('*')
        
        for file_path in files:
            if file_path.is_file() and file_path.suffix.lower() in self.SUPPORTED_EXTENSIONS:
                try:
                    doc = self.process_file(str(file_path))
                    documents.append(doc)
                except Exception as e:
                    logger.warning(f"处理文件失败 {file_path}: {e}")
        
        logger.info(f"目录处理完成: {dir_path}, 成功解析 {len(documents)} 个文件")
        return documents
    
    def _parse_pdf(self, path: Path) -> str:
        """
        解析PDF文件
        
        Args:
            path: PDF文件路径
            
        Returns:
            提取的文本内容
        """
        try:
            reader = PdfReader(str(path))
            text_parts = []
            
            for i, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text.strip():
                    text_parts.append(f"[第{i+1}页]\n{page_text}")
            
            return '\n\n'.join(text_parts)
            
        except Exception as e:
            logger.error(f"PDF解析失败 {path}: {e}")
            raise
    
    def _parse_docx(self, path: Path) -> str:
        """
        解析Word文档
        
        Args:
            path: Word文件路径
            
        Returns:
            提取的文本内容
        """
        try:
            doc = Document(str(path))
            paragraphs = []
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)
            
            # 提取表格内容
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        table_text.append(row_text)
                if table_text:
                    paragraphs.append("[表格]\n" + '\n'.join(table_text))
            
            return '\n\n'.join(paragraphs)
            
        except Exception as e:
            logger.error(f"Word解析失败 {path}: {e}")
            raise
    
    def _parse_text(self, path: Path) -> str:
        """
        解析纯文本文件 (包括.md和.txt)
        
        Args:
            path: 文本文件路径
            
        Returns:
            文件内容
        """
        try:
            with open(path, 'r', encoding=self.encoding) as f:
                return f.read()
        except UnicodeDecodeError:
            # 尝试其他编码
            for encoding in ['gbk', 'gb2312', 'utf-16']:
                try:
                    with open(path, 'r', encoding=encoding) as f:
                        logger.info(f"使用 {encoding} 编码读取文件: {path}")
                        return f.read()
                except UnicodeDecodeError:
                    continue
            raise ValueError(f"无法识别文件编码: {path}")
        except Exception as e:
            logger.error(f"文本解析失败 {path}: {e}")
            raise


def get_sample_documents() -> List[Dict[str, Any]]:
    """
    获取示例文档数据 (用于测试)
    
    Returns:
        示例文档列表
    """
    return [
        {
            "content": "规划选址应当符合城乡规划要求，避开地质灾害易发区、洪涝灾害危险区、生态保护红线等禁止建设区域。",
            "metadata": {"source": "test_regulation_1.txt", "type": "regulation"},
        },
        {
            "content": "项目建设用地应当节约集约用地，严格控制用地规模，提高土地利用效率。",
            "metadata": {"source": "test_regulation_2.txt", "type": "regulation"},
        },
        {
            "content": "选址论证报告应当包括项目概况、选址分析、合规性分析、合理性分析、节地分析、结论与建议等章节。",
            "metadata": {"source": "test_guideline_1.txt", "type": "guideline"},
        },
    ]


if __name__ == "__main__":
    # 简单测试
    processor = DocumentProcessor()
    print("支持的文件格式:", processor.SUPPORTED_EXTENSIONS)