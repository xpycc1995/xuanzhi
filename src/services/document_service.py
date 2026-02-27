"""
Word文档生成服务

使用python-docx库生成符合标准模板格式的Word文档。
"""

import os
import re
from typing import Dict, Any, Optional, List, TYPE_CHECKING
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from src.utils.logger import logger

if TYPE_CHECKING:
    from src.models.site_selection_data import SiteSelectionData


class DocumentService:
    """
    Word文档生成服务

    负责加载模板、填充内容、生成最终Word文档。
    """

    def __init__(self, template_path: Optional[str] = None):
        """
        初始化文档服务

        Args:
            template_path: Word模板文件路径,默认为templates/word_templates/标准模板.docx
        """
        if template_path is None:
            # 获取项目根目录
            project_root = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
            template_path = os.path.join(
                project_root,
                "templates",
                "word_templates",
                "标准模板.docx"
            )

        self.template_path = template_path
        logger.info(f"Word模板路径: {template_path}")

        # 样式映射表
        self.style_map = {
            "title": {"font_size": 22, "bold": True, "align": "center"},
            "heading1": {"font_size": 16, "bold": True, "color": (0, 0, 0)},
            "heading2": {"font_size": 14, "bold": True, "color": (0, 0, 0)},
            "heading3": {"font_size": 12, "bold": True},
            "normal": {"font_size": 12, "line_spacing": 1.5},
        }

    def generate_report(
        self,
        project_data: Dict[str, Any],
        chapters: Dict[str, str],
        output_path: Optional[str] = None
    ) -> str:
        """
        生成完整的Word报告

        Args:
            project_data: 项目基本信息(用于填充封面)
            chapters: 章节内容字典,格式为 {"章节号": "内容(Markdown)"}
            output_path: 输出文件路径,如果为None则自动生成

        Returns:
            生成的Word文档路径
        """
        logger.info("开始生成Word报告...")

        # 1. 加载模板
        try:
            doc = Document(self.template_path)
            logger.info("✓ 模板加载成功")
        except Exception as e:
            logger.error(f"模板加载失败: {str(e)}")
            raise

        # 2. 填充封面信息
        self._fill_cover_page(doc, project_data)

        # 3. 填充章节内容
        for chapter_num, content in chapters.items():
            self._replace_chapter_content(doc, chapter_num, content)

        # 4. 保存文档
        if output_path is None:
            output_dir = "output/reports"
            os.makedirs(output_dir, exist_ok=True)

            project_name = project_data.get("项目名称", "未命名项目")
            output_path = os.path.join(output_dir, f"{project_name}_规划选址论证报告.docx")

        try:
            doc.save(output_path)
            logger.info(f"✓ 报告生成成功: {output_path}")
            return output_path
        except Exception as e:
            logger.error(f"保存文档失败: {str(e)}")
            raise

    def _fill_cover_page(self, doc: Document, project_data: Dict[str, Any]):
        """
        填充封面信息

        Args:
            doc: Word文档对象
            project_data: 项目数据
        """
        logger.info("填充封面信息...")

        # 封面常见字段映射
        field_mapping = {
            "项目名称": ["项目名称", "【项目名称】", "{项目名称}"],
            "委托单位": ["委托单位", "建设单位", "【填写委托单位名称】"],
            "编制单位": ["编制单位", "【填写编制单位名称】"],
            "编制日期": ["编制日期", "【填写编制日期】"],
        }

        # 遍历文档前10段(通常封面在前10段)
        for paragraph in doc.paragraphs[:15]:
            for field_name, placeholders in field_mapping.items():
                # 如果项目数据中有这个字段
                if field_name in project_data:
                    value = str(project_data[field_name])

                    # 替换所有可能的占位符
                    for placeholder in placeholders:
                        if placeholder in paragraph.text:
                            paragraph.text = paragraph.text.replace(placeholder, value)
                            # 应用标题样式(检查样式是否存在)
                            if field_name == "项目名称":
                                try:
                                    paragraph.style = "Title"
                                except KeyError:
                                    # 如果Title样式不存在,尝试使用Heading 1
                                    try:
                                        paragraph.style = "Heading 1"
                                    except KeyError:
                                        # 如果都不存在,保持原样式
                                        pass

        logger.info("✓ 封面信息填充完成")

    def _replace_chapter_content(self, doc: Document, chapter_num: str, content: str):
        """
        替换指定章节的内容

        Args:
            doc: Word文档对象
            chapter_num: 章节编号(如"1", "2")
            content: 新的章节内容(Markdown格式)
        """
        logger.info(f"替换第{chapter_num}章内容...")

        # 查找章节标题（跳过目录）
        chapter_found = False
        chapter_index = -1

        for i, paragraph in enumerate(doc.paragraphs):
            # 跳过目录条目（样式为 toc 1, toc 2 等）
            style_name = paragraph.style.name.lower()
            if 'toc' in style_name:
                continue

            # 查找章节标题(如"1 项目概况"或"# 1 项目概况")
            text = paragraph.text.strip()
            if text.startswith(f"{chapter_num} ") or text.startswith(f"# {chapter_num} "):
                chapter_found = True
                chapter_index = i
                logger.info(f"找到第{chapter_num}章标题: {text} (样式: {paragraph.style.name})")
                break

        if not chapter_found:
            logger.warning(f"未找到第{chapter_num}章标题,将追加到文档末尾")
            self._append_chapter(doc, chapter_num, content)
            return

        # 找到下一章的起始位置(用于确定删除范围)
        next_chapter_index = len(doc.paragraphs)
        for i in range(chapter_index + 1, len(doc.paragraphs)):
            paragraph = doc.paragraphs[i]
            text = paragraph.text.strip()

            # 跳过目录条目和空段落
            style_name = paragraph.style.name.lower()
            if not text or 'toc' in style_name:
                continue

            # 检查是否是其他章节标题(如"2 xxx", "# 2 xxx")
            import re
            match = re.match(r'^#?\s*(\d+)\s+', text)
            if match and match.group(1) != chapter_num:
                next_chapter_index = i
                logger.info(f"找到下一章(第{match.group(1)}章)在第{i}段")
                break

        # 删除章节原有的内容(保留标题,删除后面的段落)
        # 注意:需要从后往前删除,避免索引变化
        paragraphs_to_delete = []
        for i in range(chapter_index + 1, next_chapter_index):
            paragraphs_to_delete.append(i)

        # 从后往前删除段落
        for i in reversed(paragraphs_to_delete):
            # 获取段落的父元素
            p_element = doc.paragraphs[i]._element
            p_element.getparent().remove(p_element)

        logger.info(f"删除了{len(paragraphs_to_delete)}个旧段落")

        # 在章节标题后插入新内容
        # 注意: insert_paragraph_before() 会在段落前插入，所以我们反向插入（从后往前）

        # 解析Markdown内容
        lines = content.split('\n')
        # 过滤空行并反转列表
        non_empty_lines = [line.strip() for line in lines if line.strip()]
        reversed_lines = list(reversed(non_empty_lines))

        # 获取插入参考段落
        if chapter_index + 1 < len(doc.paragraphs):
            # 获取下一章的标题段落作为参考点
            ref_paragraph = doc.paragraphs[chapter_index + 1]
        else:
            # 如果没有下一章，在文档末尾添加一个临时段落作为参考
            ref_paragraph = doc.add_paragraph()

        inserted_count = 0

        for line in reversed_lines:
            # 在参考段落前插入（这样最后插入的内容会在最前面）
            new_para = ref_paragraph.insert_paragraph_before(line)

            # 设置内容和样式
            # 判断标题级别
            if line.startswith('### '):
                new_para.text = line[4:]
                new_para.style = "Heading 3"
            elif line.startswith('## '):
                new_para.text = line[3:]
                new_para.style = "Heading 2"
            elif line.startswith('# '):
                new_para.text = line[2:]
                new_para.style = "Heading 1"
            else:
                new_para.text = line
                # 设置中文字体
                self._set_chinese_font(new_para, '宋体', 12)

            inserted_count += 1
            # 更新参考段落为当前插入的段落
            ref_paragraph = new_para

        logger.info(f"✓ 第{chapter_num}章内容替换完成 (插入了{inserted_count}行内容)")

    def _append_chapter(self, doc: Document, chapter_num: str, content: str):
        """
        追加新章节到文档末尾

        Args:
            doc: Word文档对象
            chapter_num: 章节编号
            content: 章节内容
        """
        # 添加章节标题
        title = doc.add_paragraph(f"# {chapter_num} 项目概况")
        title.style = "Heading 1"

        # 添加内容(MVP版本:简化处理)
        # 将Markdown内容按行分割,创建段落
        lines = content.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue

            # 判断标题级别
            if line.startswith('### '):
                p = doc.add_paragraph(line[4:], style="Heading 3")
            elif line.startswith('## '):
                p = doc.add_paragraph(line[3:], style="Heading 2")
            elif line.startswith('# '):
                p = doc.add_paragraph(line[2:], style="Heading 1")
            else:
                # 普通段落
                p = doc.add_paragraph(line)
                # 设置中文字体
                self._set_chinese_font(p, '宋体', 12)

    def _set_chinese_font(self, paragraph, font_name: str, font_size: int):
        """
        设置段落的中文字体

        Args:
            paragraph: 段落对象
            font_name: 字体名称
            font_size: 字体大小
        """
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)
            # 设置中文字体
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

    # ==================== 表格生成方法 ====================

    def _create_table(
        self,
        doc: Document,
        headers: List[str],
        rows: List[List[str]],
        title: str = ""
    ) -> Any:
        """
        创建专业表格

        Args:
            doc: Word文档对象
            headers: 表头列表
            rows: 数据行列表
            title: 表格标题（可选）

        Returns:
            创建的表格对象
        """
        # 创建表格
        table = doc.add_table(rows=len(rows) + 1, cols=len(headers))

        # 尝试设置表格样式，如果不存在则跳过
        try:
            table.style = 'Table Grid'
        except KeyError:
            # 如果样式不存在，手动设置边框
            pass

        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 设置表头
        header_row = table.rows[0]
        for i, header in enumerate(headers):
            cell = header_row.cells[i]
            cell.text = header
            # 表头加粗和居中
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(10.5)
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        # 填充数据行
        for row_idx, row_data in enumerate(rows):
            row = table.rows[row_idx + 1]
            for col_idx, cell_text in enumerate(row_data):
                cell = row.cells[col_idx]
                cell.text = str(cell_text)
                # 设置单元格字体
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.size = Pt(10.5)
                        run.font.name = '宋体'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        # 添加表格标题
        if title:
            title_para = doc.add_paragraph()
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = title_para.add_run(title)
            run.font.size = Pt(10.5)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        return table

    def _create_land_use_table(
        self,
        doc: Document,
        site_data: 'SiteSelectionData'
    ) -> Any:
        """
        创建表2-1：项目备选方案土地利用现状表

        Args:
            doc: Word文档对象
            site_data: 选址分析数据

        Returns:
            创建的表格对象
        """
        # 获取两个方案的土地利用数据
        scheme1 = site_data.备选方案[0]
        scheme2 = site_data.备选方案[1]

        land_use1 = scheme1.土地利用现状
        land_use2 = scheme2.土地利用现状

        # 合并所有用地类型
        all_land_types = sorted(set(land_use1.keys()) | set(land_use2.keys()))

        # 构建表格数据
        headers = ["用地类型", f"方案一（平方米）", f"方案二（平方米）"]
        rows = []

        total1 = 0.0
        total2 = 0.0

        for land_type in all_land_types:
            area1 = land_use1.get(land_type, "0")
            area2 = land_use2.get(land_type, "0")

            # 转换为数值计算合计
            try:
                val1 = float(area1.replace("平方米", "").strip()) if isinstance(area1, str) else float(area1)
            except (ValueError, AttributeError):
                val1 = 0.0

            try:
                val2 = float(area2.replace("平方米", "").strip()) if isinstance(area2, str) else float(area2)
            except (ValueError, AttributeError):
                val2 = 0.0

            total1 += val1
            total2 += val2

            rows.append([land_type, f"{val1:.2f}" if val1 > 0 else "0", f"{val2:.2f}" if val2 > 0 else "0"])

        # 添加合计行
        rows.append(["合计", f"{total1:.2f}", f"{total2:.2f}"])

        # 创建表格
        table = self._create_table(doc, headers, rows, "表2-1 项目备选方案土地利用现状表")

        # 添加数据来源说明
        source_para = doc.add_paragraph()
        source_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if site_data.数据来源:
            run = source_para.add_run(f"数据来源：{site_data.数据来源}")
        else:
            run = source_para.add_run("数据来源：2023年国土变更调查数据、勘测定界数据")
        run.font.size = Pt(9)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        return table

    def _create_investment_table(
        self,
        doc: Document,
        site_data: 'SiteSelectionData'
    ) -> Any:
        """
        创建表2-2：工程总投资情况表

        Args:
            doc: Word文档对象
            site_data: 选址分析数据

        Returns:
            创建的表格对象
        """
        construction = site_data.施工运营条件

        # 构建表格数据
        headers = ["方案", "总投资（万元）", "建安工程费", "工程建设其他费", "预备费"]
        rows = [
            ["方案一", construction.方案一总投资, "待补充", "待补充", "待补充"],
            ["方案二", construction.方案二总投资, "待补充", "待补充", "待补充"]
        ]

        # 创建表格
        table = self._create_table(doc, headers, rows, "表2-2 工程总投资情况表")

        return table

    def _create_comparison_table(
        self,
        doc: Document,
        site_data: 'SiteSelectionData'
    ) -> Any:
        """
        创建表2-3：备选方案主要指标对比表

        Args:
            doc: Word文档对象
            site_data: 选址分析数据

        Returns:
            创建的表格对象
        """
        scheme1 = site_data.备选方案[0]
        scheme2 = site_data.备选方案[1]
        external = site_data.外部配套条件
        construction = site_data.施工运营条件

        # 构建对比数据
        rows = [
            ["面积（平方米）", f"{scheme1.面积:.2f}", f"{scheme2.面积:.2f}"],
            ["占用耕地", "不占用" if not scheme1.是否占用耕地 else "占用",
             "不占用" if not scheme2.是否占用耕地 else "占用"],
            ["占用基本农田", "不占用" if not scheme1.是否占用永久基本农田 else "占用",
             "不占用" if not scheme2.是否占用永久基本农田 else "占用"],
            ["交通条件", external.交通, external.交通],
            ["外部配套", "完善", "完善"],
            ["施工难度", construction.施工难度, construction.施工难度],
            ["总投资（万元）", construction.方案一总投资, construction.方案二总投资],
            ["生态影响", "较小", "较小"],
            ["政府支持", construction.政府支持, construction.政府支持],
            ["综合评价", "优" if site_data.方案比选.推荐方案 == "方案一" else "良",
             "优" if site_data.方案比选.推荐方案 == "方案二" else "良"]
        ]

        headers = ["比选因子", "方案一", "方案二"]

        # 创建表格
        table = self._create_table(doc, headers, rows, "表2-3 备选方案主要指标对比表")

        return table

    # ==================== 第2章生成方法 ====================

    def generate_chapter_2(
        self,
        content: str,
        site_data: 'SiteSelectionData',
        output_path: Optional[str] = None
    ) -> str:
        """
        生成第2章：建设项目选址可行性分析

        包含文本内容、专业表格和图表占位符

        Args:
            content: LLM生成的Markdown内容
            site_data: 选址分析数据
            output_path: 输出文件路径

        Returns:
            生成的Word文档路径
        """
        logger.info("开始生成第2章...")

        # 加载模板
        try:
            doc = Document(self.template_path)
            logger.info("✓ 模板加载成功")
        except Exception as e:
            logger.error(f"模板加载失败: {str(e)}")
            raise

        # 替换第2章内容
        self._replace_chapter_content_enhanced(doc, "2", content, site_data)

        # 保存文档
        if output_path is None:
            output_dir = "output/reports"
            os.makedirs(output_dir, exist_ok=True)
            project_name = site_data.项目基本信息.get("项目名称", "未命名项目")
            output_path = os.path.join(output_dir, f"{project_name}_第2章.docx")

        try:
            doc.save(output_path)
            logger.info(f"✓ 第2章生成成功: {output_path}")
            return output_path
        except Exception as e:
            logger.error(f"保存文档失败: {str(e)}")
            raise

    def _replace_chapter_content_enhanced(
        self,
        doc: Document,
        chapter_num: str,
        content: str,
        site_data: 'SiteSelectionData'
    ):
        """
        替换章节内容（增强版，支持表格）

        Args:
            doc: Word文档对象
            chapter_num: 章节编号
            content: Markdown内容
            site_data: 选址数据（用于生成表格）
        """
        logger.info(f"替换第{chapter_num}章内容（增强版）...")

        # 查找章节标题
        chapter_found = False
        chapter_index = -1

        for i, paragraph in enumerate(doc.paragraphs):
            style_name = paragraph.style.name.lower()
            if 'toc' in style_name:
                continue

            text = paragraph.text.strip()
            if text.startswith(f"{chapter_num} ") or text.startswith(f"# {chapter_num} "):
                chapter_found = True
                chapter_index = i
                logger.info(f"找到第{chapter_num}章标题: {text}")
                break

        if not chapter_found:
            logger.warning(f"未找到第{chapter_num}章标题")
            return

        # 找到下一章位置
        next_chapter_index = len(doc.paragraphs)
        for i in range(chapter_index + 1, len(doc.paragraphs)):
            paragraph = doc.paragraphs[i]
            text = paragraph.text.strip()
            style_name = paragraph.style.name.lower()
            if not text or 'toc' in style_name:
                continue
            match = re.match(r'^#?\s*(\d+)\s+', text)
            if match and match.group(1) != chapter_num:
                next_chapter_index = i
                break

        # 删除原有内容
        paragraphs_to_delete = list(range(chapter_index + 1, next_chapter_index))
        for i in reversed(paragraphs_to_delete):
            p_element = doc.paragraphs[i]._element
            p_element.getparent().remove(p_element)

        logger.info(f"删除了{len(paragraphs_to_delete)}个旧段落")

        # 解析内容并插入
        # 使用正则表达式识别表格占位符
        table_patterns = {
            "表2-1": r'表2-1[^表]*',
            "表2-2": r'表2-2[^表]*',
            "表2-3": r'表2-3[^表]*'
        }

        # 在章节标题后插入内容
        # 获取插入参考点
        if chapter_index + 1 < len(doc.paragraphs):
            ref_paragraph = doc.paragraphs[chapter_index + 1]
        else:
            ref_paragraph = doc.add_paragraph()

        # 处理内容，识别表格标记位置
        lines = content.split('\n')

        # 简化处理：先插入文本，然后在末尾添加表格
        # （实际项目中可以更精细地控制表格位置）

        inserted_count = 0
        tables_inserted = {"表2-1": False, "表2-2": False, "表2-3": False}

        for line in reversed(lines):
            line = line.strip()
            if not line:
                continue

            # 检查是否是表格标记
            is_table_marker = False
            for table_name in ["表2-1", "表2-2", "表2-3"]:
                if line.startswith(table_name) and not tables_inserted[table_name]:
                    is_table_marker = True
                    break

            if not is_table_marker:
                # 插入普通文本
                new_para = ref_paragraph.insert_paragraph_before(line)

                if line.startswith('#### '):
                    new_para.text = line[5:]
                    # 尝试设置Heading 4样式，如果不存在则使用Heading 3
                    try:
                        new_para.style = "Heading 4"
                    except KeyError:
                        try:
                            new_para.style = "Heading 3"
                        except KeyError:
                            # 如果都不存在，手动设置粗体
                            for run in new_para.runs:
                                run.font.bold = True
                elif line.startswith('### '):
                    new_para.text = line[4:]
                    try:
                        new_para.style = "Heading 3"
                    except KeyError:
                        try:
                            new_para.style = "Heading 2"
                        except KeyError:
                            for run in new_para.runs:
                                run.font.bold = True
                elif line.startswith('## '):
                    new_para.text = line[3:]
                    try:
                        new_para.style = "Heading 2"
                    except KeyError:
                        try:
                            new_para.style = "Heading 1"
                        except KeyError:
                            for run in new_para.runs:
                                run.font.bold = True
                elif line.startswith('# '):
                    new_para.text = line[2:]
                    try:
                        new_para.style = "Heading 1"
                    except KeyError:
                        for run in new_para.runs:
                            run.font.bold = True
                            run.font.size = Pt(16)
                else:
                    new_para.text = line
                    self._set_chinese_font(new_para, '宋体', 12)

                inserted_count += 1
                ref_paragraph = new_para

        logger.info(f"✓ 插入了{inserted_count}行内容")

        # 在章节末尾添加表格（在下一章标题之前）
        # 由于我们之前已经删除了内容，需要在适当位置添加表格
        # 这里简化处理：在文档中找到合适的位置

        # 添加空行
        doc.add_paragraph()

        # 添加表2-1：项目备选方案土地利用现状表
        logger.info("生成表2-1...")
        self._create_land_use_table(doc, site_data)
        tables_inserted["表2-1"] = True

        # 添加空行
        doc.add_paragraph()

        # 添加表2-2：工程总投资情况表
        logger.info("生成表2-2...")
        self._create_investment_table(doc, site_data)
        tables_inserted["表2-2"] = True

        # 添加空行
        doc.add_paragraph()

        # 添加表2-3：备选方案主要指标对比表
        logger.info("生成表2-3...")
        self._create_comparison_table(doc, site_data)
        tables_inserted["表2-3"] = True

        logger.info(f"✓ 第{chapter_num}章内容替换完成，已生成3个表格")


# 测试代码
if __name__ == "__main__":
    # 测试文档服务
    print("测试Word文档生成服务...")

    try:
        # 创建文档服务(不需要真实模板)
        service = DocumentService()

        # 测试数据
        test_project_data = {
            "项目名称": "测试项目",
            "委托单位": "测试单位",
            "编制单位": "AI开发团队",
            "编制日期": "2025年2月"
        }

        test_chapters = {
            "1": "## 项目背景\n\n这是测试内容..."
        }

        print("\n✓ 文档服务初始化成功")
        print("注意:实际生成需要标准模板文件")

    except Exception as e:
        print(f"\n✗ 测试失败: {str(e)}")
